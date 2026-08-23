"""Render a Document Generator block-model into a true-RTL .docx.

Applies real bidirectional formatting, not just right-alignment:
  - paragraphs get `<w:bidi/>` + right alignment + `<w:rtl/>` runs
  - tables get `<w:bidiVisual/>` so column order mirrors automatically
  - code/diagram blocks are forced LTR + monospace so ASCII art stays intact
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

_CODE_FONT = "Consolas"
_BASE_FONT = "Arial"


def _bidi(pPr, value: str = "true") -> None:
    existing = pPr.find(qn("w:bidi"))
    if existing is None:
        existing = OxmlElement("w:bidi")
        pPr.append(existing)
    existing.set(qn("w:val"), value)


def _rtl_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _bidi(paragraph._p.get_or_add_pPr(), "true")
    for run in paragraph.runs:
        run._r.get_or_add_rPr().append(OxmlElement("w:rtl"))


def _ltr_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _bidi(paragraph._p.get_or_add_pPr(), "false")


def _rtl_table(table) -> None:
    tblPr = table._tbl.tblPr
    if tblPr.find(qn("w:bidiVisual")) is None:
        tblPr.append(OxmlElement("w:bidiVisual"))


def _configure_document(doc) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = _BASE_FONT
    normal.font.size = Pt(11)
    section_pr = doc.sections[0]._sectPr
    if section_pr.find(qn("w:bidi")) is None:
        section_pr.append(OxmlElement("w:bidi"))


def _fill_table(table, rows: list[list[str]], *, bold_header: bool) -> None:
    """Fill by table geometry, padding short rows and ignoring extra cells so a
    malformed model row cannot IndexError the whole render (LLM-output boundary)."""
    table.style = "Table Grid"
    _rtl_table(table)
    cols = len(table.columns)
    for r, row in enumerate(rows):
        for c in range(cols):
            value = row[c] if c < len(row) else ""
            cell = table.cell(r, c)
            cell.text = str(value)
            paragraph = cell.paragraphs[0]
            if bold_header and r == 0:
                for run in paragraph.runs:
                    run.bold = True
            _rtl_paragraph(paragraph)


def _render_heading(doc, block: dict) -> None:
    try:
        level = int(block.get("level", 1))
    except (TypeError, ValueError):
        level = 1
    level = max(0, min(level, 9))
    heading = doc.add_heading(block.get("text", ""), level=level)
    _rtl_paragraph(heading)


def _render_paragraph(doc, block: dict) -> None:
    _rtl_paragraph(doc.add_paragraph(block.get("text", "")))


def _render_field(doc, block: dict) -> None:
    paragraph = doc.add_paragraph()
    label = paragraph.add_run(block.get("label", ""))
    label.bold = True
    paragraph.add_run(" " + block.get("value", ""))
    _rtl_paragraph(paragraph)


def _render_bullets(doc, block: dict) -> None:
    for item in block.get("items", []):
        _rtl_paragraph(doc.add_paragraph(str(item), style="List Bullet"))


def _render_field_table(doc, block: dict) -> None:
    rows = block.get("rows", [])
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=2)
    _fill_table(table, rows, bold_header=False)


def _render_table(doc, block: dict) -> None:
    headers = block.get("headers", [])
    rows = block.get("rows", [])
    if not headers:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    _fill_table(table, [headers, *rows], bold_header=True)


def _render_code(doc, block: dict) -> None:
    paragraph = doc.add_paragraph()
    for i, line in enumerate(block.get("text", "").split("\n")):
        if i:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.name = _CODE_FONT
        run.font.size = Pt(9)
    _ltr_paragraph(paragraph)


def _render_page_break(doc, _block: dict) -> None:
    doc.add_page_break()


_RENDERERS = {
    "heading": _render_heading,
    "paragraph": _render_paragraph,
    "field": _render_field,
    "bullets": _render_bullets,
    "field_table": _render_field_table,
    "table": _render_table,
    "code": _render_code,
    "page_break": _render_page_break,
}


def render_document(model: dict, output_path: str) -> str:
    """Render the block-model to a true-RTL .docx and return the output path."""
    doc = Document()
    _configure_document(doc)

    title = model.get("title")
    if title:
        _rtl_paragraph(doc.add_heading(title, level=0))

    for block in model.get("blocks", []):
        renderer = _RENDERERS.get(block.get("type"))
        if renderer is None:
            raise ValueError(f"unknown block type: {block.get('type')!r}")
        renderer(doc, block)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def _main() -> None:
    import json
    import sys

    if len(sys.argv) != 3:
        raise SystemExit("usage: python -m src.renderer <model.json> <output.docx>")
    model = json.loads(Path(sys.argv[1]).read_text(encoding="utf-8"))
    print(render_document(model, sys.argv[2]))


if __name__ == "__main__":
    _main()
